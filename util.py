import os

DOC_OUTPUT = 'target/output.docx'
TREE_CHILD = '|---'
TREE_CHILD2 = '|   '


def get_docx(file_path):
    from docx import Document
    doc_file = None
    if os.path.exists(file_path):
        doc_file = Document(file_path)
    else:
        doc_file = Document()
    return doc_file


def dump_source_code(file_path, doc_obj=None):
    """
    把源代码打印到word文档里去
    :param file_path: 源代码路径
    :param doc_obj: .docx文档对象
    :return:
    """
    from docx.shared import Pt, Cm
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y-%m-%d::')
    doc_obj.add_paragraph(timestamp + 'file_path:: ' + file_path)

    with open(file_path, 'r') as file:
        line_num = 0
        file_text = ''
        for line in file:
            line_num += 1
            line_num_str = str(line_num)
            while len(line_num_str) < 3:
                line_num_str = '0' + line_num_str
            file_text += line_num_str+' '*4+line
        para = doc_obj.add_paragraph(file_text)
        para_format = para.paragraph_format
        para_format.space_before = Cm(0)
        para_format.space_after = Cm(0)
        para_format.line_spacing = Pt(12)



def tree_dir(file_path, depth=0):
    # 先打印根目录
    if depth == 0:
        print(file_path)
    depth += 1
    files = os.listdir(file_path)
    for item in files:
        child_str = TREE_CHILD2*(depth-1)+TREE_CHILD + item
        print(child_str)
        # os.path.isdir 一定要传入完整路径, 否则单给文件夹名字, 是无法判断是不是文件夹的
        item_path = os.path.join(file_path, item)
        if  os.path.isdir(item_path):
            tree_dir(item_path, depth)


class TreeDir:
    str_trees = ''

    def __init__(self, file_path):
        self.tree_dir_str(file_path)

    def tree_dir_str(self, file_path, depth=0):
        # 先打印根目录
        if depth == 0:
            self.str_trees += (file_path + '\n')
        depth += 1
        files = os.listdir(file_path)
        for item in files:
            child_str = TREE_CHILD2*(depth-1)+TREE_CHILD + item
            self.str_trees += (child_str + '\n')
            # os.path.isdir 一定要传入完整路径, 否则单给文件夹名字, 是无法判断是不是文件夹的
            item_path = os.path.join(file_path, item)
            if  os.path.isdir(item_path):
                self.tree_dir_str(item_path, depth)


def traverse_dump(file_path, doc_obj=None):
    files = os.listdir(file_path)
    for item in files:
        # os.path.isdir 一定要传入完整路径, 否则单给文件夹名字, 是无法判断是不是文件夹的
        item_path = os.path.join(file_path, item)
        if  os.path.isdir(item_path):
            traverse_dump(item_path,doc_obj)
        else:
            dump_source_code(item_path, doc_obj)

if __name__ == '__main__':
    a = TreeDir('source/extensions').str_trees
    print(a)
