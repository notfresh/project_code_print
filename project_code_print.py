# coding:utf-8
import os
import sys


TREE_CHILD = '|---'
TREE_CHILD2 = '|   '

include_suffix = ['.py', '.html']
exclude_prefix = ['.', 'test_', '__pycache__', 'project_code_print.py']

exclude_suffix = ['.pyc', '.docx', 'pdf']


def get_docx(file_path):
    """
    获取docx文件
    :param file_path:
    :return:
    """
    from docx import Document
    doc_file = None
    if os.path.exists(file_path):
        doc_file = Document(file_path)
    else:
        doc_file = Document()
    return doc_file



class TreeDir:

    def __init__(self, file_path):
        self.str_trees = ''
        self.tree_dir_str(file_path)

    def tree_dir_str(self, file_path, depth=0):
        """
        打印给定目录的树形结构
        :param file_path:
        :param depth:
        :return:
        """
        # 先打印根目录
        file_path_last = file_path.split(os.path.sep)[-1]
        if depth == 0:
            self.str_trees += (file_path_last + '\n')

        depth += 1
        files = os.listdir(file_path)
        for item in files:
            # 排除特定前缀, 后缀
            flag = True
            for v in exclude_prefix:
                if item.startswith(v):
                    flag = False
                    break
            for v in exclude_suffix:
                if item.endswith(v):
                    flag = False
                    break
            if not flag:
                continue
            child_str = TREE_CHILD2*(depth-1) + TREE_CHILD + item
            self.str_trees += (child_str + '\n')
            # os.path.isdir 一定要传入完整路径, 否则单给文件夹名字, 是无法判断是不是文件夹的
            item_path = os.path.join(file_path, item)
            if os.path.isdir(item_path):
                self.tree_dir_str(item_path, depth)


def dump_source_code(file_path, doc_obj):
    """
    把源代码打印到word文档里去
    :param file_path: 源代码路径
    :param doc_obj: .docx文档对象
    :return:
    """
    # 排除特定后缀
    flag = False
    for v in include_suffix:
        if file_path.endswith(v):
            flag = True
            break
    if not flag:
        return

    from docx.shared import Pt, Cm
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y-%m-%d::')


    title = doc_obj.add_paragraph(timestamp + 'file_path:: ' + file_path)
    para_format = title.paragraph_format
    para_format.space_before = Pt(12)

    with open(file_path, 'r', encoding='utf-8') as file:
        line_num = 0
        file_text = ''
        try:
            for line in file:
                line_num += 1
                line_num_str = str(line_num)
                while len(line_num_str) < 3:
                    line_num_str = '0' + line_num_str
                file_text += line_num_str+' '*4+line
        except Exception:
            print('出现异常')
            print(file_path)
            print(line_num)
        para = doc_obj.add_paragraph(file_text)
        para_format = para.paragraph_format
        para_format.space_before = Cm(0)
        para_format.space_after = Cm(0)
        para_format.line_spacing = Pt(12)


def traverse_dump(file_path, doc_obj):
    files = os.listdir(file_path)
    for item in files:
        # 排除特定前缀
        flag = True
        for v in exclude_prefix:
            if item.startswith(v):
                flag = False
                break
        if not flag:
            continue

        # os.path.isdir 一定要传入完整路径, 否则单给文件夹名字, 是无法判断是不是文件夹的
        item_path = os.path.join(file_path, item)
        if os.path.isdir(item_path):
            traverse_dump(item_path,doc_obj)
        else:
            dump_source_code(item_path, doc_obj)

try:
    import docx
except Exception:
    print("当前python的运行环境没有安装 python-docx 模块, 请使用 pip install python-docx 进行安装.")
    sys.exit(0)

# 自定义排除的
exclude_prefix += ['本项目-', ]

if __name__ == '__main__':
    dump_dir_path = os.getcwd()
    output_file = 'code_print.docx'
    doc_obj = get_docx(output_file)
    # 添加文件目录树
    doc_obj.add_paragraph(TreeDir(dump_dir_path).str_trees)
    # 添加源代码
    traverse_dump(dump_dir_path, doc_obj)
    doc_obj.save(output_file)

