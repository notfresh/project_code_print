from docx import Document
from docx.shared import Inches, Pt

doc = Document()

def demo1():
    doc.add_heading('This is title', 0)
    p = doc.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some')
    p.add_run(' italic.').italic = True
    doc.add_heading('Heading level 1', level=1)
    doc.add_paragraph('Intense quote', style='IntenseQuote')
    doc.add_paragraph('first item in unordered list', style='ListBullet')
    doc.add_paragraph('first item in ordered list', style='ListNumber')
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].txt = 'Qty'
    hdr_cells[1].txt = 'Id'
    hdr_cells[2].txt = 'Desc'
    for item in table.rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.qty)
        row_cells[1].text = str(item.id)
        row_cells[2].text = item.desc
    doc.add_page_break()

def demo2():
    # p0 = doc.add_paragraph('Normal text0 ')
    # p0.add_run('text with emphasis', 'Emphasis')
    # pf = p0.paragraph_format
    # pf.left_indent = Inches(0)
    #
    # p = doc.add_paragraph('Normal text ')
    # p.add_run('text with emphasis', 'Emphasis')
    # pf = p.paragraph_format
    # pf.left_indent = Inches(0.5)
    #
    # p2 = doc.add_paragraph('Normal text2 ')
    # p2.add_run('text with emphasis', 'Emphasis')
    # pf2 = p2.paragraph_format
    # pf2.left_indent = Inches(1)
    #
    # p3 = doc.add_paragraph('Normal text3')
    # p3.add_run('text with emphasis', 'Emphasis')
    # pf = p3.paragraph_format
    # pf.first_line_indent = Inches(-0.25)  #  伸出去了

    # p4 = doc.add_paragraph('Normal text4')
    # pf = p4.paragraph_format
    # pf.space_before = Pt(0)
    # pf.space_after = Pt(0)
    #
    # p5 = doc.add_paragraph('Normal text5')
    # pf = p5.paragraph_format
    # pf.space_before = Pt(0)
    # pf.space_after = Pt(0)
    #
    # p6 = doc.add_paragraph('Normal text6')
    # pf = p6.paragraph_format
    # pf.space_before = Pt(0)
    # pf.space_after = Pt(3)

    for i in range(0,50):
        p = doc.add_paragraph(str(i)+'   Normal text6')
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

def demo3():
    p = doc.add_paragraph('''
        Line spacing is the distance between subsequent baselines in the lines of a paragraph. Line spacing can be specified either as an absolute distance or relative to the line height (essentially the point size of the font used). A typical absolute measure would be 18 points. A typical relative measure would be double-spaced (2.0 line heights). The default line spacing is single-spaced (1.0 line heights).

    Line spacing is controlled by the interaction of the line_spacing and line_spacing_rule properties. line_spacing is either a Length value, a (small-ish) float, or None. A Length value indicates an absolute distance. A float indicates a number of line heights. None indicates line spacing is inherited. line_spacing_rule is a member of the WD_LINE_SPACING enumeration or None:

        ''')

    p = doc.add_paragraph('''
    Line spacing is the distance between subsequent baselines in the lines of a paragraph. Line spacing can be specified either as an absolute distance or relative to the line height (essentially the point size of the font used). A typical absolute measure would be 18 points. A typical relative measure would be double-spaced (2.0 line heights). The default line spacing is single-spaced (1.0 line heights).

Line spacing is controlled by the interaction of the line_spacing and line_spacing_rule properties. line_spacing is either a Length value, a (small-ish) float, or None. A Length value indicates an absolute distance. A float indicates a number of line heights. None indicates line spacing is inherited. line_spacing_rule is a member of the WD_LINE_SPACING enumeration or None:
    
    ''')
    pf = p.paragraph_format
    pf.line_spacing = Pt(18)

    p = doc.add_paragraph('''
        Line spacing is the distance between subsequent baselines in the lines of a paragraph. Line spacing can be specified either as an absolute distance or relative to the line height (essentially the point size of the font used). A typical absolute measure would be 18 points. A typical relative measure would be double-spaced (2.0 line heights). The default line spacing is single-spaced (1.0 line heights).

    Line spacing is controlled by the interaction of the line_spacing and line_spacing_rule properties. line_spacing is either a Length value, a (small-ish) float, or None. A Length value indicates an absolute distance. A float indicates a number of line heights. None indicates line spacing is inherited. line_spacing_rule is a member of the WD_LINE_SPACING enumeration or None:

        ''')
    pf = p.paragraph_format
    pf.line_spacing = Pt(10)

    p = doc.add_paragraph('''
            Line spacing is the distance between subsequent baselines in the lines of a paragraph. Line spacing can be specified either as an absolute distance or relative to the line height (essentially the point size of the font used). A typical absolute measure would be 18 points. A typical relative measure would be double-spaced (2.0 line heights). The default line spacing is single-spaced (1.0 line heights).

        Line spacing is controlled by the interaction of the line_spacing and line_spacing_rule properties. line_spacing is either a Length value, a (small-ish) float, or None. A Length value indicates an absolute distance. A float indicates a number of line heights. None indicates line spacing is inherited. line_spacing_rule is a member of the WD_LINE_SPACING enumeration or None:

            ''')
    pf = p.paragraph_format
    pf.line_spacing = Pt(12)


    p = doc.add_paragraph('''
            Line spacing is the distance between subsequent baselines in the lines of a paragraph. Line spacing can be specified either as an absolute distance or relative to the line height (essentially the point size of the font used). A typical absolute measure would be 18 points. A typical relative measure would be double-spaced (2.0 line heights). The default line spacing is single-spaced (1.0 line heights).

        Line spacing is controlled by the interaction of the line_spacing and line_spacing_rule properties. line_spacing is either a Length value, a (small-ish) float, or None. A Length value indicates an absolute distance. A float indicates a number of line heights. None indicates line spacing is inherited. line_spacing_rule is a member of the WD_LINE_SPACING enumeration or None:

            ''')
    pf = p.paragraph_format
    pf.line_spacing = 1



demo3()
doc.save('demo.docx')
